from pathlib import Path
import re
from typing import Any

EMPTY_CONTENT_ERROR = "文件内容为空，无法建立知识库。"
PARSE_FAILED_ERROR = "文件解析失败，请确认文件未损坏或更换格式后重试。"
SCANNED_PDF_ERROR = "该 PDF 可能是扫描件或不含可提取文本，当前版本暂不支持 OCR。"


class DocumentParseError(Exception):
    def __init__(self, message: str) -> None:
        self.message = message
        super().__init__(message)


def parse_document(path: Path, original_filename: str, extension: str) -> str:
    try:
        if extension == ".pdf":
            return parse_pdf(path, original_filename)
        if extension == ".docx":
            return parse_docx(path, original_filename)
        if extension == ".xlsx":
            return parse_xlsx(path, original_filename)
        if extension in {".txt", ".md"}:
            return parse_text_file(path, original_filename)
    except DocumentParseError:
        raise
    except Exception as exc:
        raise DocumentParseError(PARSE_FAILED_ERROR) from exc

    raise DocumentParseError("文件类型不支持，请上传 PDF、DOCX、XLSX、TXT 或 Markdown 文件。")


def parse_pdf(path: Path, original_filename: str) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise DocumentParseError(PARSE_FAILED_ERROR) from exc

    try:
        reader = PdfReader(str(path))
    except Exception as exc:
        raise DocumentParseError(PARSE_FAILED_ERROR) from exc

    sections: list[str] = []
    for page_index, page in enumerate(reader.pages, start=1):
        page_text = normalize_extracted_text(page.extract_text() or "")
        if page_text:
            sections.append(f"来源文件：{original_filename}\n页码：{page_index}\n正文：{page_text}")

    if not sections:
        raise DocumentParseError(SCANNED_PDF_ERROR)
    return "\n\n".join(sections)


def parse_docx(path: Path, original_filename: str) -> str:
    try:
        from docx import Document as DocxDocument
    except ImportError as exc:
        raise DocumentParseError(PARSE_FAILED_ERROR) from exc

    document = DocxDocument(str(path))
    sections: list[str] = []

    paragraphs = [normalize_extracted_text(paragraph.text) for paragraph in document.paragraphs]
    paragraph_text = "\n".join(paragraph for paragraph in paragraphs if paragraph)
    if paragraph_text:
        sections.append(f"来源文件：{original_filename}\n正文：{paragraph_text}")

    for table_index, table in enumerate(document.tables, start=1):
        rows = [
            [normalize_extracted_text(cell.text) for cell in row.cells]
            for row in table.rows
        ]
        table_lines = table_rows_to_text(rows)
        if table_lines:
            sections.append(f"来源文件：{original_filename}\n表格：{table_index}\n" + "\n".join(table_lines))

    if not sections:
        raise DocumentParseError(EMPTY_CONTENT_ERROR)
    return "\n\n".join(sections)


def parse_xlsx(path: Path, original_filename: str) -> str:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise DocumentParseError(PARSE_FAILED_ERROR) from exc

    workbook = load_workbook(filename=str(path), read_only=True, data_only=True)
    sections: list[str] = []

    for sheet in workbook.worksheets:
        if getattr(sheet, "sheet_state", "visible") != "visible":
            continue

        rows: list[list[str]] = []
        for row in sheet.iter_rows(values_only=True):
            values = [cell_to_text(value) for value in row]
            if any(values):
                rows.append(values)

        sheet_lines = table_rows_to_text(rows)
        if sheet_lines:
            sections.append(f"来源文件：{original_filename}\nSheet：{sheet.title}\n" + "\n".join(sheet_lines))

    workbook.close()
    if not sections:
        raise DocumentParseError(EMPTY_CONTENT_ERROR)
    return "\n\n".join(sections)


def parse_text_file(path: Path, original_filename: str) -> str:
    raw = path.read_bytes()
    decoded_text: str | None = None
    for encoding in ("utf-8", "utf-8-sig"):
        try:
            decoded_text = raw.decode(encoding)
            break
        except UnicodeDecodeError:
            continue

    if decoded_text is None:
        raise DocumentParseError(PARSE_FAILED_ERROR)

    cleaned = normalize_markdown_text(decoded_text if path.suffix.lower() == ".md" else decoded_text)
    if not cleaned:
        raise DocumentParseError(EMPTY_CONTENT_ERROR)
    return f"来源文件：{original_filename}\n正文：{cleaned}"


def table_rows_to_text(rows: list[list[str]]) -> list[str]:
    rows = [[cell for cell in row] for row in rows if any(row)]
    if not rows:
        return []

    headers = rows[0]
    data_rows = rows[1:] if len(rows) > 1 else rows
    use_headers = len(rows) > 1 and any(headers)
    lines: list[str] = []

    for row in data_rows:
        pairs: list[str] = []
        for index, value in enumerate(row):
            if not value:
                continue
            if use_headers and index < len(headers) and headers[index]:
                pairs.append(f"{headers[index]}：{value}")
            else:
                pairs.append(value)
        if pairs:
            lines.append("；".join(pairs))

    return lines


def cell_to_text(value: Any) -> str:
    if value is None:
        return ""
    return normalize_extracted_text(str(value))


def normalize_extracted_text(text: str) -> str:
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def normalize_markdown_text(text: str) -> str:
    text = re.sub(r"```[\s\S]*?```", lambda match: match.group(0).replace("```", ""), text)
    text = re.sub(r"^\s{0,3}#{1,6}\s*", "", text, flags=re.MULTILINE)
    text = re.sub(r"[*_`~]{1,3}", "", text)
    return normalize_extracted_text(text)
